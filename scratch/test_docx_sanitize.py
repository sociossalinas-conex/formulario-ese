import zipfile
import io
import os
import docx

def sanitize_docx_zip(filepath):
    temp_io = io.BytesIO()
    dummy_png = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15c4\x00\x00\x00\nIDATx\x9cc`\x00\x00\x00\x02\x00\x01H\xaf\xa4q\x00\x00\x00\x00IEND\xaeB`\x82'
    
    with zipfile.ZipFile(filepath, 'r') as zin:
        with zipfile.ZipFile(temp_io, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                try:
                    data = zin.read(item.filename)
                    zout.writestr(item, data)
                except Exception as e:
                    print(f"    [Sanitizing] Replacing corrupt file: {item.filename} ({e})")
                    if item.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                        new_item = zipfile.ZipInfo(item.filename)
                        new_item.compress_type = zipfile.ZIP_DEFLATED
                        zout.writestr(new_item, dummy_png)
                    else:
                        new_item = zipfile.ZipInfo(item.filename)
                        new_item.compress_type = zipfile.ZIP_DEFLATED
                        zout.writestr(new_item, b'')
                        
    with open(filepath, 'wb') as f:
        f.write(temp_io.getvalue())

template_dir = r"c:\Users\anton\OneDrive\Documentos\Antigravity\Formulario Captura Final\Plantillas"
filepath = os.path.join(template_dir, "Afoso.docx")

print("1. Sanitizing Afoso.docx zip...")
sanitize_docx_zip(filepath)

print("2. Attempting to load with python-docx...")
try:
    doc = docx.Document(filepath)
    print("Success loading Afoso.docx with python-docx!")
    print(f"Paragraphs found: {len(doc.paragraphs)}")
    print(f"Tables found: {len(doc.tables)}")
except Exception as e:
    print("Failed loading after sanitization:", e)
