import docx
import os
import re

template_dir = r"c:\Users\anton\OneDrive\Documentos\Antigravity\Formulario Captura Final\Plantillas"
files = ["Lease.docx", "Festo.docx", "Transcar.docx", "Afoso.docx"]

for fname in files:
    path = os.path.join(template_dir, fname)
    if not os.path.exists(path):
        print(f"File {fname} not found!")
        continue
    
    print(f"\n=================== INSPECTING {fname} ===================")
    doc = docx.Document(path)
    
    # Let's count paragraphs and tables
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Search for anything that looks like a placeholder: [ ... ], {{ ... }}, or underscores
    placeholders = []
    
    # Inspect paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        # Find brackets or braces or multiple underscores
        p_matches = re.findall(r"\[.*?\]|\{\{.*?\}\}|_{3,}", txt)
        if p_matches:
            placeholders.append((p.text, p_matches))
            
    # Inspect tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if not txt:
                    continue
                c_matches = re.findall(r"\[.*?\]|\{\{.*?\}\}|_{3,}", txt)
                if c_matches:
                    placeholders.append((cell.text, c_matches))
                    
    print(f"Found {len(placeholders)} matches. Examples:")
    for text, match in placeholders[:15]:
        print(f"  Text: {text!r} -> Matches: {match}")
