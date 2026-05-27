import docx
import os
import re

def to_snake_case(text):
    # Strip formatting
    text = re.sub(r"\*\*|\*|__|_|#|`|:", "", text)
    text = text.strip()
    replacements = {
        'á': 'a', 'é': 'e', 'í': 'i', 'ó': 'o', 'ú': 'u',
        'Á': 'a', 'É': 'e', 'Í': 'i', 'Ó': 'o', 'Ú': 'u',
        'ñ': 'n', 'Ñ': 'n',
        'ü': 'u', 'Ü': 'u',
        '?': '', '¿': '', '(': '', ')': '', '/': '_', '-': '_', '.': '', ',': ''
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"[\s_]+", "_", text)
    text = text.lower().strip("_")
    
    mappings = {
        'nombre_completo': 'nombre',
        'nombre_del_candidato': 'nombre',
        'nombre_del_prospecto': 'nombre',
        'fecha_de_visita': 'fecha_visita',
        'fecha_de_visita_domiciliaria': 'fecha_visita',
        'fecha_nacimiento': 'fecha_nacimiento',
        'fecha_de_nacimiento': 'fecha_nacimiento',
        'lugar_nacimiento': 'lugar_nacimiento',
        'lugar_de_nacimiento': 'lugar_nacimiento',
        'estado_de_nacimiento': 'lugar_nacimiento',
        'estado_nacimiento': 'lugar_nacimiento',
        'genero': 'genero',
        'sexo': 'genero',
        'estado_civil': 'estado_civil',
        'curp': 'curp',
        'rfc': 'rfc',
        'edad': 'edad',
        'telefono_de_casa': 'telefono_casa',
        'telefono_casa': 'telefono_casa',
        'telefono_de_recados': 'telefono_recados',
        'telefono_recados': 'telefono_recados',
        'telefono_para_recados': 'telefono_recados',
        'celular': 'celular',
        'telefono_celular': 'celular',
        'correo': 'correo',
        'correo_electronico': 'correo',
        'email': 'correo',
        'tipo_de_sangre': 'tipo_sangre',
        'tipo_sangre': 'tipo_sangre',
        'grupo_sanguineo': 'tipo_sangre',
        'licencia': 'licencia',
        'tiene_licencia': 'licencia',
        'tipo_de_licencia': 'tipo_licencia',
        'tipo_licencia': 'tipo_licencia',
        'vencimiento': 'vencimiento_licencia',
        'vencimiento_de_licencia': 'vencimiento_licencia',
        'vencimiento_licencia': 'vencimiento_licencia',
        'contacto_de_emergencia': 'contacto_emergencia',
        'nombre_emergencia': 'emergencia_nombre',
        'emergencia_nombre': 'emergencia_nombre',
        'telefono_emergencia': 'emergencia_telefono',
        'emergencia_telefono': 'emergencia_telefono',
        'nombre_del_padre': 'nombre_padre',
        'nombre_padre': 'nombre_padre',
        'nombre_de_la_madre': 'nombre_madre',
        'nombre_madre': 'nombre_madre'
    }
    
    if text in mappings:
        return mappings[text]
    return text

def update_docx_placeholders(filepath, output_path):
    doc = docx.Document(filepath)
    
    # 1. Update paragraphs outside tables (e.g. Label: ________)
    for p in doc.paragraphs:
        txt = p.text
        if not txt.strip():
            continue
            
        # Match e.g., **Label**: ______ or Label: ______ or Label: [ ]
        # Regex: matches a label ending in colon followed by underscores or empty brackets
        pattern = r"([a-zA-ZáéíóúÁÉÍÓÚñÑ0-9\s\-/().]{2,40})\s*:\s*(_{3,}|\[\s*\])"
        matches = re.findall(pattern, txt)
        if matches:
            new_text = txt
            for label, fill in matches:
                snake_id = to_snake_case(label)
                if snake_id and len(snake_id) > 1:
                    new_text = new_text.replace(f"{label}:{fill}", f"{label}: {{{{{snake_id}}}}}")
                    new_text = new_text.replace(f"{label}: {fill}", f"{label}: {{{{{snake_id}}}}}")
                    # In case of spaces or formatting
                    new_text = re.sub(re.escape(label) + r"\s*:\s*" + re.escape(fill), f"{label}: {{{{{snake_id}}}}}", new_text)
            p.text = new_text

    # 2. Update table cells
    for table in doc.tables:
        for row in table.rows:
            # We check pairs of cells
            for i in range(len(row.cells) - 1):
                cell_a = row.cells[i]
                cell_b = row.cells[i+1]
                
                txt_a = cell_a.text.strip()
                txt_b = cell_b.text.strip()
                
                clean_lbl = re.sub(r"\*\*|\*|__|_|#|`|:", "", txt_a).strip()
                is_valid_label = (
                    len(clean_lbl) > 1 and 
                    len(clean_lbl) < 60 and 
                    not clean_lbl.startswith("---") and
                    not txt_b.startswith("{{") and
                    not txt_a.startswith("{{")
                )
                
                is_empty_or_fill = (
                    txt_b == "" or 
                    re.match(r"^_{2,}$", txt_b) or
                    txt_b == "[ ]" or
                    txt_b == "[]"
                )
                
                if is_valid_label and is_empty_or_fill:
                    snake_id = to_snake_case(txt_a)
                    if snake_id and len(snake_id) > 1:
                        # Set text in cell B, preserving any styling if possible by editing paragraphs
                        if cell_b.paragraphs:
                            cell_b.paragraphs[0].text = f"{{{{{snake_id}}}}}"
                        else:
                            cell_b.text = f"{{{{{snake_id}}}}}"
                            
    doc.save(output_path)

template_dir = r"c:\Users\anton\OneDrive\Documentos\Antigravity\Formulario Captura Final\Plantillas"
filepath = os.path.join(template_dir, "Lease.docx")
out_path = os.path.join(template_dir, "../scratch/Lease_modified.docx")

update_docx_placeholders(filepath, out_path)
print("Updated Lease.docx saved in scratch!")
